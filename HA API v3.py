import os
import re
import time
import logging
import PyPDF2
import docx
import requests
from tqdm import tqdm
from tenacity import retry, stop_after_attempt, wait_exponential
from PIL import Image
from io import BytesIO

# Configuration
MAX_TOKENS_PER_CHUNK = 2000
QUESTION_TYPES = {
    1: "Case-based diagnosis from radiological features",
    4: "Disease characteristics verification",
    6: "Special feature identification"
}
HF_API_TOKEN = " "  #add your HF API token here
HF_API_URL = "https://api-inference.huggingface.co/models/mistralai/Mistral-7B-Instruct-v0.3"
DEFAULT_OUTPUT_DIR = "C:\\Mine\\python practice\\Generated stuff"

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

def extract_text_from_pdf(pdf_path, start_page, end_page):
    """Extract text from specified PDF pages with validation"""
    text = ""
    try:
        with open(pdf_path, 'rb') as pdf_file:
            reader = PyPDF2.PdfReader(pdf_file)
            total_pages = len(reader.pages)

            if start_page < 1 or end_page > total_pages or start_page > end_page:
                raise ValueError(f"Invalid page range (1-{total_pages})")

            logging.info(f"Processing pages {start_page}-{end_page} of {total_pages}")

            for page_num in range(start_page - 1, end_page):
                page = reader.pages[page_num]
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n\n"

            return text
    except Exception as e:
        logging.error(f"PDF extraction failed: {e}")
        raise

@retry(stop=stop_after_attempt(3), wait=wait_exponential(multiplier=1, min=2, max=10))
def generate_question_with_hf(chunk, question_type):
    """Generate question using Hugging Face API"""
    prompt = f"""Generate a challenging {QUESTION_TYPES[question_type]} question using this exact format:

[CONTENT]
{chunk[:1500]}

[REQUIREMENTS]
1. Question must be based on the content
2. Include four distinct options
3. Mark one correct answer
4. Use advanced medical terminology
5. Make the question difficult

[FORMAT]
Question: [Your question here]
a) [Option A]
b) [Option B]
c) [Option C]
d) [Option D]
Correct Answer: [Letter only]"""

    headers = {"Authorization": f"Bearer {HF_API_TOKEN}"}

    try:
        response = requests.post(
            HF_API_URL,
            headers=headers,
            json={
                "inputs": prompt,
                "parameters": {
                    "max_new_tokens": 300,
                    "temperature": 0.7,
                    "top_p": 0.9
                }
            }
        )

        if response.status_code == 200:
            return response.json()[0]['generated_text'].split(prompt)[-1].strip()
        else:
            logging.error(f"API Error: {response.status_code} - {response.text}")
            return None

    except Exception as e:
        logging.error(f"Generation failed: {str(e)}")
        return None

def validate_api_token():
    """Validate the provided API token"""
    headers = {"Authorization": f"Bearer {HF_API_TOKEN}"}
    try:
        response = requests.post(
            HF_API_URL,
            headers=headers,
            json={"inputs": "test"}
        )
        if response.status_code != 200:
            logging.error(f"API token validation failed: {response.status_code} - {response.text}")
            raise ValueError("Invalid API token or URL")
    except Exception as e:
        logging.error(f"API token validation failed: {str(e)}")
        raise

def validate_question(text):
    """Validate the generated question format"""
    try:
        if not text or len(text.strip()) < 50:
            return False

        required_elements = [
            r'Question:',
            r'a\)', r'b\)', r'c\)', r'd\)',
            r'Correct Answer: [a-d]'
        ]

        return all(re.search(pattern, text, re.IGNORECASE)
                   for pattern in required_elements)
    except Exception as e:
        logging.error(f"Validation error: {str(e)}")
        return False

def save_to_word(content, q_type, doc_path, chunk_id):
    """Save generated questions to Word document"""
    try:
        doc = docx.Document(doc_path) if os.path.exists(doc_path) else docx.Document()
        if not doc.paragraphs:
            doc.add_heading("Generated Medical Questions", 0)
            doc.add_paragraph(f"Generated on {time.strftime('%Y-%m-%d %H:%M:%S')}\n")

        doc.add_heading(f"Section {chunk_id}: {QUESTION_TYPES[q_type]}", level=2)
        doc.add_paragraph(content)
        doc.save(doc_path)
        return True
    except Exception as e:
        logging.error(f"Document save failed: {str(e)}")
        return False

def process_text_chunk(chunk, doc_path, chunk_id):
    """Process a single chunk of text"""
    successful = 0
    for q_type in QUESTION_TYPES:
        result = generate_question_with_hf(chunk, q_type)
        if result and validate_question(result):
            if save_to_word(result, q_type, doc_path, chunk_id):
                successful += 1
    return successful

def main():
    try:
        print("Medical Question Generator")
        print("==========================")

        # Validate API token
        validate_api_token()

        # User inputs
        pdf_path = input("\nEnter full path to PDF file: ").strip()
        page_range = input("Enter page range (e.g., 1-10): ").strip()
        start_page, end_page = map(int, page_range.split('-'))
        output_file = input("Output file name (e.g., questions.docx): ").strip()

        # Ensure the output file has a .docx extension
        if not output_file.endswith(".docx"):
            output_file += ".docx"

        # Ensure the output directory exists
        if not os.path.exists(DEFAULT_OUTPUT_DIR):
            os.makedirs(DEFAULT_OUTPUT_DIR)

        # Full path for the output file
        full_output_path = os.path.join(DEFAULT_OUTPUT_DIR, output_file)

        # Process document
        print("\nExtracting text from PDF...")
        text_content = extract_text_from_pdf(pdf_path, start_page, end_page)

        # Split into chunks (simple paragraph-based splitting)
        chunks = [text_content[i:i+MAX_TOKENS_PER_CHUNK]
                  for i in range(0, len(text_content), MAX_TOKENS_PER_CHUNK)]

        print(f"\nGenerating questions for {len(chunks)} text chunks...")
        total_valid = 0

        with tqdm(total=len(chunks) * len(QUESTION_TYPES),
                  desc="Generating Questions",
                  unit="question") as pbar:
            for idx, chunk in enumerate(chunks):
                valid = process_text_chunk(chunk, full_output_path, idx + 1)
                total_valid += valid
                pbar.update(len(QUESTION_TYPES))

        print(f"\nSuccessfully created {total_valid} questions in {full_output_path}")
        print(f"File saved to: {full_output_path}")
        print("Note: The Hugging Face API has rate limits. For heavy usage, consider:")
        print("- Using a paid inference endpoint")
        print("- Running a local model instead")

    except Exception as e:
        logging.error(f"\nFatal error: {str(e)}")
        print("Please check:")
        print("- Your API token is valid")
        print("- PDF path is correct")
        print("- Page numbers are valid")

if __name__ == "__main__":
    main()